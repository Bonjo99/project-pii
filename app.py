from flask import *
from flask import jsonify #solo per errore too many request
from flask import render_template, send_file
from flask import flash
from flask import redirect, url_for, request, session, make_response
from flask_dance.contrib.google import make_google_blueprint, google
import shutil, os, datetime, random
from azure.storage.blob import BlobServiceClient, BlobClient, ContainerClient
from azure.ai.textanalytics import TextAnalyticsClient
from azure.identity import DefaultAzureCredential
from azure.core.credentials import AzureKeyCredential
from azure.core.exceptions import HttpResponseError
from werkzeug.utils import secure_filename
import tempfile
import io 
from reportlab.pdfgen import canvas
from PyPDF2 import PdfReader, PdfWriter
import pdfplumber
from docx import Document
import docx
from flask import request
from flask_bcrypt import Bcrypt
from password_strength import PasswordPolicy
from zxcvbn import zxcvbn
import pwnedpasswords
import mysql.connector
from mysql.connector import Error
import re
import os
from argon2 import PasswordHasher
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import logging
from concurrent.futures import ThreadPoolExecutor
import base64
from requests_oauthlib import OAuth2Session
import json
import fitz
from flask_mail import Mail, Message
import string
from markupsafe import escape

app = Flask(__name__, static_folder="static", template_folder="template")
app.secret_key = os.urandom(24)


#config google
app.config["GOOGLE_OAUTH_CLIENT_ID"] = "1086144218901-66r02mo1suk7qdibb6cijtgkrmrr8a9j.apps.googleusercontent.com"  # Google Client ID
app.config["GOOGLE_OAUTH_CLIENT_SECRET"] = "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"  #Google Client Secret
google_bp = make_google_blueprint(
    client_id=app.config["GOOGLE_OAUTH_CLIENT_ID"],
    client_secret=app.config["GOOGLE_OAUTH_CLIENT_SECRET"],
    scope=["https://www.googleapis.com/auth/userinfo.email", "https://www.googleapis.com/auth/userinfo.profile", "openid"],
    redirect_to="google_login"  # Specifies the endpoint where you handle the OAuth callback
)

app.register_blueprint(google_bp, url_prefix="/login")
os.environ['OAUTHLIB_INSECURE_TRANSPORT'] = '1'
bcrypt = Bcrypt(app)
language_key = "523f45988b474bba8e9f474b8e97c4b0"
language_endpoint = "https://piilanguagesistemi.cognitiveservices.azure.com/"
user="piiadmin2024"
db_password="ProgettoSistemi2024"
blob_service_client = BlobServiceClient.from_connection_string('DefaultEndpointsProtocol=https;AccountName=piiprofilearchive;AccountKey=hsZDEcnQwHUAbqjD6moAbSmb0XeC6LXfuanPq9LT+ZFf/csxE/NdyQxKXSUNhKkWVytmQKAt142v+AStDJLWBw==;EndpointSuffix=core.windows.net')
policy = PasswordPolicy.from_names(length
                                   =12,)

def rf(a):
	return open(a,"r").read()

# Authenticate the client using your key and endpoint 
def authenticate_client():
    ta_credential = AzureKeyCredential(language_key)
    text_analytics_client = TextAnalyticsClient(
            endpoint=language_endpoint, 
            credential=ta_credential)
    return text_analytics_client

client = authenticate_client()

# Configure Flask-Mail
app.config['MAIL_SERVER'] = 'smtp.googlemail.com'
app.config['MAIL_PORT'] = 587
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USERNAME'] = 'francesco.bongiovanni99@gmail.com'
app.config['MAIL_PASSWORD'] = 'usac qqra ymbi idmk'
mail = Mail(app)


# Configure the session to use the filesystem (you can also configure it to use databases or redis)
app.config['SESSION_TYPE'] = 'filesystem'

@app.after_request
def no_cache(response):
    """
    Add no caching headers.
    """
    response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
    response.headers["Pragma"] = "no-cache"
    response.headers["Expires"] = "0"
    response.headers['Cache-Control'] = 'public, max-age=0'
    return response
# Database Connection
def create_conn():
    conn = mysql.connector.connect(host='piidatabaseserver.mysql.database.azure.com',
                                   database='pii',
                                   user=user,
                                   password=db_password)
    return conn

executor = ThreadPoolExecutor(max_workers=2)

@app.route('/')
def index():
    return render_template('index.html')

@app.route("/dashboard", methods=["GET"])
def dashboard():
    if 'username' in session:
        username = session['username']
        container_client = blob_service_client.get_container_client(username)
        files = [{"name": blob.name, "size": str(round(blob.size / 1024 / 1024)) + " mb"} for blob in container_client.list_blobs()]
        space = sum([blob.size for blob in container_client.list_blobs()]) / 1024 / 1024
        return render_template("dashboard.html", name=username, files=files, username=username, space=round(space), nf=len(files))
    else:
        return redirect(url_for('sign_in'))

#GOOGLE LOGIN
@app.route("/google_login")
def google_login():
    if not google.authorized:
        return redirect(url_for("google.login"))

    response = google.get("/oauth2/v2/userinfo")
    if response.ok:
        google_data = response.json()
        email = google_data["email"]
        name = google_data["name"]

        # Check if the user exists
        try:
            conn = create_conn()
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM Users WHERE email = %s", (email,))
            account = cursor.fetchone()

            if account:
                # User exists, log them in
                session['username'] = account[0]  # Assuming the first column is the username
                session['email'] = account[2]  # Assuming the third column is the email
                return redirect(url_for("dashboard"))
            else:
                # User does not exist, redirect them to choose a username
                session['email'] = email
                session['name'] = name
                return redirect(url_for("choose_username"))

        except Exception as e:
            return str(e)  # You should handle the error more gracefully in production

    return "Error accessing Google data"

@app.route("/choose_username", methods=["GET", "POST"])
def choose_username():
    if request.method == 'POST':
        username = request.form["username"].strip()
        # Check if the username is taken
        try:
            conn = create_conn()
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM Users WHERE username = %s", (username,))
            account = cursor.fetchone()

            if account:
                # Username is taken, show an error message
                flash("Username is already taken", "error")
            if not is_valid_username(username):
                print("Username must contain only lowercase letters and numbers")
                flash("Username must contain only lowercase letters and numbers", "error")
                return render_template("choose_username.html")
            else:
                # Username is available, create a new user
                password = bcrypt.generate_password_hash(os.urandom(24)).decode('utf-8')
                cursor.execute("INSERT INTO Users (username, password, email, nome_utente) VALUES (%s, %s, %s, %s)",
                               (username, password, session['email'], session['name']))
                conn.commit()
                cursor.close()
                conn.close()

                # Create a new container for the user
                blob_service_client.create_container(username)

                session['username'] = username
                return redirect(url_for('dashboard', username=username))

        except Exception as e:
            return str(e)  # You should handle the error more gracefully in production

    return render_template("choose_username.html")


@app.route("/reset_password", methods=["GET", "POST"])
def reset_password():
    if request.method == "POST":
        email = str(escape(request.form["email"]))
        with create_conn() as conn:
            with conn.cursor() as cursor:
                cursor.execute("SELECT * FROM Users WHERE email = %s", (email,))
                account = cursor.fetchone()
                if account:
                    # Generate a new random password
                    new_password = ''.join(random.choices(string.ascii_uppercase + string.digits, k=12))
                    hashed_password = bcrypt.generate_password_hash(new_password).decode('utf-8')
                    # Update the user's password in the database.
                    cursor.execute("UPDATE Users SET password = %s WHERE email = %s", (hashed_password, email))
                    conn.commit()
                    # Send the new password to the user via email.
                    msg = Message('Your password has been reset',
                                  sender='noreply@demo.com',
                                  recipients=[email])
                    msg.body = f'Your new password is: {new_password}'
                    mail.send(msg)
                    return render_template("signin.html", message="Your password has been reset. Please check your email.")
                else:
                    return render_template("reset_password.html", message="No account with that email address exists.")
    return render_template("reset_password.html", message="")

@app.route("/sign_up", methods=["GET","POST"])
def sign_up():
    form_data = session.get('form_data', {
        "username": "",
        "email": "",
        "name": ""
    })
    if request.method == 'POST':
        # Retrieve registration form data
        password = escape(request.form["password"])
        username = escape(request.form["username"])
        email = escape(request.form["email"])
        name = escape(request.form["name"])
        form_data = {
            "username": username,
            "email": email,
            "name": name
        }
        session['form_data'] = form_data
        
        # Check whether the password has been previously used asynchronously using a ThreadPoolExecutor to avoid blocking the main thread
        future = executor.submit(check_pwned_password, password)

        # Check the length of the password
        if policy.test(password):
            flash("Password must be at least 12 characters long", "error")
        # Verify password strength
        if zxcvbn(password)['score'] < 3:
            flash("Password is too weak", "error")
        # Checks whether the password has been used before.
        # Verifies whether the username consists of only lowercase letters and numbers.
        if not is_valid_username(username):
            print("Username must contain only lowercase letters and numbers")
            flash("Username must contain only lowercase letters and numbers", "error")
            return render_template("signup.html", message="Username must contain only lowercase letters and numbers", form_data=form_data)
        
        # Retrieve registration form data
        user_data = {
            "username": request.form["username"],
            "password": bcrypt.generate_password_hash(request.form["password"]).decode('utf-8'),
            "email": request.form["email"],
            "name": request.form["name"]
        }
        # Check if the user or email already exists.
        try:
            print("Checking if user or email exists")
            with create_conn() as conn:
                with conn.cursor() as cursor:
                    cursor.execute("SELECT * FROM Users WHERE username = %s OR email = %s", (user_data["username"], user_data["email"]))
                    account = cursor.fetchone()
                    if (future.result()):                
                        return render_template("signup.html", message="Password has been pwned", form_data=form_data)
                    if account:
                        return render_template("signup.html", message="User or email already exists", form_data=form_data)
                    else: 
                        query = "INSERT INTO Users (username, password, email, nome_utente) VALUES (%s, %s, %s, %s)"
                        print(f"Executing query: {query}")
                        cursor.execute(query, (user_data["username"], user_data["password"], user_data["email"], user_data["name"]))
                        conn.commit()
        except Exception as e:
            print(f"Error checking if user or email exists: {e}")
            return render_template("signup.html", message="An error occurred: " + str(e), form_data=form_data)
            
        # Crea un contenitore per l'utente
        try:
            blob_service_client.create_container(user_data["username"])
            session['username'] = user_data["username"]
            session.pop('_flashes', None)
            #return render_template("dashboard.html", files=[], password=user_data["password"], username=username, space=0, nf=0)
            return redirect(url_for('dashboard', username=username))
        except Exception as e:
            print(f"Error creating container for user: {e}")
            return render_template("signup.html", message="An error occurred: " + str(e))
    return render_template("signup.html", message="", form_data=form_data)

def is_valid_username(username):
    return all(c.islower() or c.isdigit() for c in username)
def check_pwned_password(password):
    return pwnedpasswords.check(password)
#LOGIN   
@app.route("/sign_in", methods=["GET","POST"])
def sign_in():
    if (request.method == "GET"):
        return render_template("signin.html", message="")
    else:
        username = str(escape(request.form["username"]))
        password = escape(request.form["password"])
        try:
            conn = create_conn()
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM Users WHERE username = %s", (username,))
            account = cursor.fetchone()
            if account:
                hashed_password = account[1]
                if bcrypt.check_password_hash(hashed_password, password):
                    # Set session and redirect to dashboard
                    session['username'] = username
                    return redirect(url_for('dashboard', username=username))
                else:
                    return render_template("signin.html", message="Wrong password")
            else:
                return render_template("signin.html", message="No such account")
        except Exception as e:
            print(f"Error checking if user exists: {e}")
            return render_template("signin.html", message="An error occurred")
        
#UPLOAD FILE
@app.route("/upload/<user>/file", methods=["POST"])
def upload(user):
    if 'username' not in session:
        return render_template("dashboard.html", message="You must be signed in to upload files")
    if request.method == "POST":
        container_client = blob_service_client.get_container_client(user)
        files = request.files.getlist("files[]")
        valid_files = []
         # Calculate the total size of the files to be uploaded
        total_size = sum(f.content_length for f in files)

        # Check if the user has enough space
        space_used = sum([blob.size for blob in container_client.list_blobs()]) / 1024 / 1024
        space_left = 1024 - space_used  # 1GB - space used
        if total_size > space_left:
            flash("Not enough space. Please delete some files or upgrade your storage.")
            return redirect(url_for('dashboard', username=user))

        # Validate files first
        for f in files:
            original_filename = f.filename
            name, extension = os.path.splitext(original_filename)
            if extension not in ['.docx', '.pdf']:
                flash("Invalid file type. Only .docx and .pdf files are allowed.")
                return redirect(url_for('dashboard', username=user))
            valid_files.append(f)

        # If validation passes, proceed with uploading
        for f in valid_files:
            original_filename = f.filename
            counter = 1
            file_blob = container_client.get_blob_client(f.filename)
            name, extension = os.path.splitext(f.filename)
            while file_blob.exists():
                new_filename = f"{name}({counter}){extension}"
                file_blob = container_client.get_blob_client(new_filename)
                counter += 1
            file_blob.upload_blob(f.read())

            document_name = file_blob.blob_name
            try:
                if extension == '.docx':
                    txt = docx_to_string(file_blob)
                elif extension == '.pdf':
                    txt = pdf_to_string(file_blob)
                documents = [txt]
                metadata, confidence_score = pii_detection(documents, user, extension, file_blob, container_client)
                metadata_json = json.dumps(metadata)
                conn = create_conn()
                cursor = conn.cursor()
                cursor.execute(
                    "INSERT INTO document (document_name, username, metadata) VALUES (%s, %s, %s)",
                    (document_name, user, metadata_json)
                )
                conn.commit()
                cursor.close()
                conn.close()
            except HttpResponseError as e:
                if e.status_code == 413:
                    flash("Document too large. Limit document size to: 5120 text elements.")
                    return redirect(url_for('dashboard', username=user,confidence_score=confidence_score))
                else:
                    raise

        files = [{"name": blob.name, "size": str(round(blob.size / 1024 / 1024)) + " mb"} for blob in container_client.list_blobs()]
        space = sum([blob.size for blob in container_client.list_blobs()]) / 1024 / 1024
        print("File uploaded con confidence",confidence_score)
        flash("Privacy Risk Level: " + confidence_score)
        return redirect(url_for('dashboard', username=user))
    else:
        return "Invalid request method", 405

#CONVERT FILE TO ANONIMIZED FILE
@app.route("/convert/<user>/name/<name>", methods=["GET"])
def convert_text_to_anonimizedtext(user, name):
    if 'username' not in session:
        return render_template("dashboard.html", message="You must be signed in to upload files")
    if request.method == "GET":
        files_container = blob_service_client.get_container_client(user)
        file_blob = files_container.get_blob_client(name)
        files = [{"name": blob.name, "size": str(round(blob.size / 1024 / 1024)) + " mb"} for blob in files_container.list_blobs()]
         # Calculate the total size of the files to be uploaded
        total_size = sum(float(f['size'].split()[0]) for f in files)

        # Check if the user has enough space
        space_used = sum([blob.size for blob in files_container.list_blobs()]) / 1024 / 1024
        space_left = 1024 - space_used  # 1GB - space used
        if total_size > space_left:
            flash("Not enough space. Please delete some files or upgrade your storage.")
            return redirect(url_for('dashboard', username=user))
        if file_blob.exists():
            name, extension = os.path.splitext(name)
            if extension == '.docx':
                txt = docx_to_string(file_blob)
            elif extension == '.pdf':
                txt = pdf_to_string(file_blob)
            if len(txt.split()) > 5120:
                flash("Error: Document size exceeds the limit of 5120 words")
                return redirect(url_for('dashboard', username=user))
            documents = [txt]
            output_file = "anonimized_" + name + extension

            # Perform PII recognition and draft the redacted document
            try:
                language_country = client.detect_language(documents, country_hint="us")[0]
                language = language_country.primary_language['iso6391_name']
            except Exception as e:
                    flash("Error: Document size exceeds the limit of 5120 words")
                    return redirect(url_for('dashboard', username=user))
            print("LINGUA", language)
            response = client.recognize_pii_entities(documents, language=language)
            
            metadata = {}
            for doc in response:
                if not doc.is_error:
                    for entity in doc.entities:
                        metadata[entity.text] = entity.category
                else:
                    print("Error:", doc.error.message)

            if metadata:
                if extension == ".pdf":
                    redact_pdf(file_blob, output_file, metadata, files_container)
                elif extension == ".docx":
                    redact_docx(file_blob, output_file, metadata, files_container)

                print("Testo convertito salvato come", output_file)
            else:
                print("No PII detected")
            files = [{"name": blob.name, "size": str(round(blob.size / 1024 / 1024)) + " mb"} for blob in files_container.list_blobs()]
            space = sum([blob.size for blob in files_container.list_blobs()]) / 1024 / 1024
            flash("Privacy Risk Level: No Risk")
            return redirect(url_for('dashboard', username=user))
    else:
        return "Invalid request method", 405

#FUNCTIONS FOR REDACTING PDF AND DOCX FILES
def redact_pdf(input_blob, output_file, metadata, container_client):
    blob_data = input_blob.download_blob().readall()
    
    # Create a temporary file to process the PDF
    with tempfile.NamedTemporaryFile(delete=False) as temp:
        temp.write(blob_data)
        temp_file_name = temp.name

    # Open the original PDF with PyMuPDF
    doc = fitz.open(temp_file_name)
    
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        for word in metadata.keys():
            text_instances = page.search_for(word)
            
            for inst in text_instances:
                # Draw a white rectangle over the text to redact
                rect = fitz.Rect(inst)
                page.draw_rect(rect, color=(1, 1, 1), fill=(1, 1, 1))
                
                # Ensure the text is completely redacted
                page.add_redact_annot(rect, fill=(1, 1, 1))  # White fill
                page.apply_redactions()

    # Save the redacted PDF
    redacted_output_path = f"{temp_file_name}_redacted.pdf"
    doc.save(redacted_output_path, encryption=fitz.PDF_ENCRYPT_KEEP)
    doc.close()
    
    # Upload the redacted PDF to the blob storage
    output_blob = container_client.get_blob_client(output_file)
    with open(redacted_output_path, "rb") as data:
        output_blob.upload_blob(data, overwrite=True)

    # Clean up temporary files
    os.remove(temp_file_name)
    os.remove(redacted_output_path)

def redact_docx(input_blob, output_file, metadata, container_client):
    blob_data = input_blob.download_blob().readall()
    output = BytesIO()
    doc = Document(io.BytesIO(blob_data))
    for paragraph in doc.paragraphs:
        redacted_paragraph = paragraph.text
        for word in metadata.keys():
            redacted_paragraph = redacted_paragraph.replace(word, '*' * len(word))
        paragraph.text = redacted_paragraph
    doc.save(output)
    output.seek(0)
    output_blob = container_client.get_blob_client(output_file)
    output_blob.upload_blob(output.read(), overwrite=True)

def docx_to_string(blob_client):
    blob_data = blob_client.download_blob().readall()
    document = Document(io.BytesIO(blob_data))
    text = ' '.join([paragraph.text for paragraph in document.paragraphs])
    return text

def pdf_to_string(blob_client):
    blob_data = blob_client.download_blob().readall()
    with pdfplumber.open(io.BytesIO(blob_data)) as pdf:
        text = ' '.join([page.extract_text() for page in pdf.pages])
    return text

def pii_detection(documents, name, extension, file_blob, files_container):
    metadata_file = "metadata_" + name + ".json"
    response = client.recognize_pii_entities(documents, language="en")
    confidence_scores = []  # List to save confidence scores
    metadata={}
    pii_found=False
    categories_to_anonymize = ["Email", "PhoneNumber", "Address"]  #  Add the categories to be anonymized here
    for doc in response:
        if not doc.is_error:
            for entity in doc.entities:
                if entity.category in categories_to_anonymize:  # Checks whether the entity category is to be anonymized
                    confidence_score = entity.confidence_score
                    confidence_scores.append(confidence_score)  # Save confidence score to list
                    metadata[entity.text] = entity.category
                    pii_found = True
        else:
            print("Error:", doc.error.message)
    if not pii_found:
        return {}, "No PII found"

    # Calculates the average confidence score and displays it on the page
    if confidence_scores:
        # Extract the Average Confidence Score from the metadata, if it exists
        average_confidence_score = sum(confidence_scores) / len(confidence_scores)
        #Confidence level
        if average_confidence_score==0:
            text="No Risk"
        if average_confidence_score < 0.25:
            text="Very Low"
            return metadata, text
        elif average_confidence_score < 0.5:
            text="Low"
            return metadata, text
        elif average_confidence_score < 0.75:
            text="Medium"
            return metadata, text
        else:
            text="High"
            return metadata, text

#DELETE FILE ON BLOB STORAGE AND DATABASE
@app.route("/delete/<user>/<name>", methods=["GET"])
def delete(user, name):
    if 'username' not in session:
        return render_template("dashboard.html", message="You must be signed in to upload files")
    print("Cancellazione file")
    if (request.method == "GET"):
                try:
                    files_container = blob_service_client.get_container_client(user)
                    file_blob = files_container.get_blob_client(name)
                    if file_blob.exists():
                        file_blob.delete_blob()
                        # If the file name does not start with "anonimized", delete it from the database
                        if not name.startswith("anonimized"):
                            try:
                                conn = create_conn()
                                cursor = conn.cursor()
                                cursor.execute("DELETE FROM document WHERE document_name = %s AND username = %s", (name, user))
                                conn.commit()
                                cursor.close()
                                conn.close()
                            except Exception as e:
                                print(f"Error deleting document from database: {e}")
                                return "An error occurred during database deletion", 500
                        files = [{"name": blob.name, "size": str(round(blob.size / 1024 / 1024)) + " mb"} for blob in files_container.list_blobs()]
                        space = sum([blob.size for blob in files_container.list_blobs()]) / 1024 / 1024
                        return redirect(url_for('dashboard', username=user))
                    else:
                        return "File does not exist", 404
                except Exception as e:
                    print(f"Error during file deletion: {e}")
                    return "An error occurred during file deletion", 500
    return "Invalid request method", 405


#DOWNLOAD FILE
@app.route("/download/<user>/name/<name>", methods=["GET"])
def download(user, name):
    if 'username' not in session:
        return redirect(url_for('index'))
    if (request.method == "GET"):
                files_container = blob_service_client.get_container_client(user)
                file_blob = files_container.get_blob_client(name)
                if file_blob.exists():
                    download_file_path = os.path.join(tempfile.gettempdir(), name)
                    with open(download_file_path, "wb") as download_file:
                        download_file.write(file_blob.download_blob().readall())
                    return send_from_directory(tempfile.gettempdir(), name, as_attachment=True)
    else:
            return "Invalid request method", 405
    
@app.route("/logout", methods=["GET"])
def logout():
    session.clear()
    return redirect(url_for('index'))

#SEARCH FILE ON BLOB STORAGE AND DATABASE(SEARCH BY NAME AND METADATA)
@app.route('/search/<user>', methods=['POST','GET'])
def search(user):
    if 'username' not in session:
        return render_template("dashboard.html", message="You must be signed in to upload files")
    
    if request.method == 'POST':
        query = escape(request.form['search'])
        if query:
            conn = create_conn()
            cursor = conn.cursor(dictionary=True)
            
            # Perform search on document name and metadata
            search_query = """
                SELECT * FROM document 
                WHERE username = %s AND 
                (document_name LIKE %s OR JSON_EXTRACT(metadata, '$') LIKE %s)
            """
            search_term = f"%{query}%"
            cursor.execute(search_query, (user, search_term, search_term))
            
            results = cursor.fetchall()
            cursor.close()
            conn.close()
            files_db = [{"name": result["document_name"]} for result in results]
            print("Files", files_db)
            
            files_container = blob_service_client.get_container_client(user)
            files_db_names = [file["name"] for file in files_db]  # Get a list of all file names in files_db
            files = []
            for blob in files_container.list_blobs():
                if blob.name in files_db_names:  # Only include the file if its name is in files_db_names
                    files.append({"name": blob.name, "size": str(round(blob.size / 1024 / 1024)) + " mb"})
                    print("Files", files)
            
            space = sum([blob.size for blob in files_container.list_blobs()]) / 1024 / 1024
            print("NOME UTENT",user)
            # Fetch user details for rendering the dashboard
            return render_template("dashboard.html", username=user, space=round(space), nf=len(files), files=files)
        else:
            return redirect(url_for('dashboard', username=user))
    else: 
        return render_template("dashboard.html", username=user, space=round(space), nf=len(files), files=files)

#VIEW FILE DIRECTLY ON BROWSER
@app.route("/view/<user>/name/<name>", methods=["GET"])
def view_file(user, name):
    print("View file")
    if 'username' not in session:
        return redirect(url_for('index'))
    if request.method == "GET":
        files_container = blob_service_client.get_container_client(user)
        file_blob = files_container.get_blob_client(name)
        if file_blob.exists():
            file_extension = os.path.splitext(name)[1].lower()
            blob_data = file_blob.download_blob().readall()
            
            if file_extension in ['.pdf', '.docx']:
                encoded_blob_data = base64.b64encode(blob_data).decode('utf-8')
                return render_template("view_file.html", username=user, file_data=encoded_blob_data, file_extension=file_extension, file_name=name)
            else:
                return "File type not supported for inline viewing", 400
        else:
            return "File not found", 404
    else:
        return "Invalid request method", 405

app.logger.setLevel(logging.DEBUG)
if __name__ == '__main__':
    app.run(debug=True)
